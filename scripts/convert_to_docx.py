#!/usr/bin/env python3
"""
convert_to_docx.py - Markdown → Word 转换
用法: python convert_to_docx.py <markdown_file> [--output DOCX_PATH]
"""

import sys
import re
import argparse
from pathlib import Path

def convert_to_docx(md_path: str, output_path: str = None) -> str:
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
    except ImportError:
        print("ERROR: 需要 python-docx 库")
        print("安装方法: pip install python-docx")
        sys.exit(1)

    md_file = Path(md_path)
    if not md_file.exists():
        print(f"ERROR: Markdown 文件不存在: {md_path}")
        sys.exit(1)

    if not output_path:
        output_path = str(md_file.with_suffix('.docx'))

    md_content = md_file.read_text(encoding='utf-8')
    doc = Document()

    # 设置默认字体（支持中文）
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    # 中文支持
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), '宋体')

    def add_heading(text: str, level: int):
        h = doc.add_heading(text, level=level)
        h.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def add_paragraph(text: str, bold: bool = False):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        return p

    def add_table_from_md(table_md: str):
        rows = [l for l in table_md.strip().split('\n') if l and '|' in l]
        if len(rows) < 2:
            return
        # 跳过表头分隔行
        data_rows = [r for r in rows if not re.match(r'\|[-| :]+\|', r)]
        headers = [h.strip() for h in data_rows[0].split('|')[1:-1]]

        table = doc.add_table(rows=len(data_rows), cols=len(headers))
        table.style = 'Table Grid'

        for i, row_data in enumerate(data_rows):
            cells = [c.strip() for c in row_data.split('|')[1:-1]]
            for j, cell_text in enumerate(cells):
                cell = table.rows[i].cells[j]
                cell.text = cell_text
                if i == 0:
                    cell.paragraphs[0].runs[0].bold = True

    lines = md_content.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].strip()

        # 跳过空行
        if not line:
            i += 1
            continue

        # 标题
        if line.startswith('# '):
            add_heading(line[2:], 0)
        elif line.startswith('## '):
            add_heading(line[3:], 1)
        elif line.startswith('### '):
            add_heading(line[4:], 2)
        elif line.startswith('#### '):
            add_heading(line[5:], 3)

        # 表格（多行）
        elif line.startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            add_table_from_md('\n'.join(table_lines))
            continue

        # 无序列表
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(re.sub(r'\*\*(.+?)\*\*', r'\1', line[2:]))

        # 任务列表
        elif re.match(r'^- \[x\] ', line):
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run('☑ ' + re.sub(r'\*\*(.+?)\*\*', r'\1', line[6:]))
        elif re.match(r'^- \[ \] ', line):
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run('☐ ' + re.sub(r'\*\*(.+?)\*\*', r'\1', line[6:]))

        # 普通段落
        else:
            # 处理 Markdown 格式
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', line)
            text = re.sub(r'\*(.+?)\*', r'\1', text)
            text = re.sub(r'`([^`]+)`', r'\1', text)
            # 处理图片
            img_match = re.match(r'!\[([^\]]*)\]\(([^\)]+)\)', text)
            if img_match:
                alt, src = img_match.groups()
                try:
                    p = doc.add_paragraph()
                    run = p.add_run(f'[{alt}] {src}')
                    run.font.color.rgb = RGBColor(0, 102, 204)
                except Exception:
                    doc.add_paragraph(text)
            else:
                doc.add_paragraph(text)

        i += 1

    doc.save(output_path)
    print(f"✅ Word 文档已生成: {output_path}")
    return output_path

if __name__ == '__main__':
    parser = argparse.ArgumentParser(description='将 Markdown 转换为 Word')
    parser.add_argument('markdown', help='Markdown 文件路径')
    parser.add_argument('--output', '-o', help='输出 Word 路径（默认与输入同名）')
    args = parser.parse_args()

    convert_to_docx(args.markdown, args.output)
